"""Експорт юридичних сторінок сайту у .docx на фірмовому бланку.

Джерелом тексту є самі Jinja-шаблони (app/templates/main/*.html), тому
документ завжди збігається з тим, що бачить користувач на сайті: сторінка
рендериться у HTML, з неї береться контентний блок і перекладається у
Word-розмітку.

Основою документа є бланк листа (docs/legal/Шаблон листа ІПРМ.docx): з нього
успадковується верхній колонтитул (логотип + реквізити) і береться зображення
печатки з підписом для підписного блоку.
"""
import io
import re
import zipfile
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from flask import current_app, render_template

# Реєстр сторінок, доступних для експорту: ключ -> (шаблон, ім'я файла)
LEGAL_PAGES = {
    'offer': ('main/offer.html', 'ІПРМ - Публічна оферта.docx'),
    'privacy': ('main/privacy.html', 'ІПРМ - Політика конфіденційності.docx'),
    'cookies': ('main/cookies.html', 'ІПРМ - Політика Cookie.docx'),
    'disclaimer': ('main/disclaimer.html', 'ІПРМ - Дисклеймер медичної відповідальності.docx'),
    'refund': ('main/refund.html', 'ІПРМ - Політика повернення коштів.docx'),
}

# Значення за замовчуванням; перевизначаються через app.config
DEFAULT_LETTERHEAD = 'docs/legal/Шаблон листа ІПРМ.docx'
DEFAULT_OUTPUT_DIR = 'docs/legal'
DEFAULT_SIGNER_TITLE = 'Ректор'
DEFAULT_SIGNER_NAME = 'Заболотня Д. О.'
SEAL_MEDIA = 'word/media/image1.png'  # печатка з підписом у бланку листа

FONT = 'Times New Roman'
BODY_PT = 12
BLACK = RGBColor(0, 0, 0)

MARGINS_CM = {'left': 2.0, 'right': 1.0, 'top': 3.0, 'bottom': 2.0}

HEADING_STYLES = {'h2': 'Heading 2', 'h3': 'Heading 3', 'h4': 'Heading 4'}
HEADING_FORMAT = {
    # стиль: (кегль, інтервал перед, інтервал після)
    'Heading 1': (16, 0, 12),
    'Heading 2': (13, 14, 6),
    'Heading 3': (12, 10, 4),
    'Heading 4': (12, 8, 4),
}


class LegalDocxError(RuntimeError):
    """Помилка експорту юридичної сторінки у .docx."""


# --- публічний API --------------------------------------------------------

def export_page(page_key, output_dir=None, with_seal=True):
    """Згенерувати .docx для сторінки з LEGAL_PAGES. Повертає шлях до файла."""
    try:
        template, filename = LEGAL_PAGES[page_key]
    except KeyError:
        raise LegalDocxError(
            'Невідома сторінка "%s". Доступні: %s'
            % (page_key, ', '.join(sorted(LEGAL_PAGES)))
        )

    soup = _render_page(template)
    content = (soup.select_one('.iprm-content-page__inner')
               or soup.select_one('.legal-page')
               or soup.find('main'))
    if content is None:
        raise LegalDocxError(
            'У шаблоні %s не знайдено контентний блок '
            '(.iprm-content-page__inner / .legal-page / <main>).' % template
        )

    doc = _open_letterhead()
    _set_page(doc)
    _set_styles(doc)
    _set_metadata(doc, content)

    _build_body(doc, content)
    if with_seal:
        _add_signature(doc)
    _add_page_number_footer(doc)

    out_dir = Path(output_dir or current_app.config.get(
        'LEGAL_DOCX_OUTPUT_DIR', DEFAULT_OUTPUT_DIR))
    if not out_dir.is_absolute():
        out_dir = _project_root() / out_dir
    out_dir.mkdir(parents=True, exist_ok=True)

    out_path = out_dir / filename
    doc.save(out_path)
    return out_path


# --- джерело контенту -----------------------------------------------------

def _project_root():
    return Path(current_app.root_path).parent


def _letterhead_path():
    path = Path(current_app.config.get('LEGAL_DOCX_LETTERHEAD', DEFAULT_LETTERHEAD))
    if not path.is_absolute():
        path = _project_root() / path
    if not path.exists():
        raise LegalDocxError('Бланк листа не знайдено: %s' % path)
    return path


def _render_page(template):
    """Відрендерити сторінку цілком і повернути дерево HTML.

    Рендеримо повний шаблон (а не вирізаний блок), щоб працювали url_for,
    переклади та контекстні процесори; потрібний блок вибираємо вже з HTML.
    """
    with current_app.test_request_context('/'):
        html = render_template(template)
    return BeautifulSoup(html, 'html.parser')


# --- каркас документа -----------------------------------------------------

def _open_letterhead():
    """Бланк як основа: лишаємо колонтитул, тіло листа очищаємо."""
    doc = Document(_letterhead_path())
    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)
    # Малюнки з тіла листа лишилися б "осиротілими" частинами пакета
    # (близько 1 МБ баласту) -- відв'язуємо їх.
    for rel_id, rel in list(doc.part.rels.items()):
        if rel.reltype.endswith('/image'):
            doc.part.drop_rel(rel_id)
    return doc


def _force_font(style, size=None, bold=None, color=None):
    """Шрифт стилю: python-docx пише лише w:ascii, кирилиця потребує решти."""
    style.font.name = FONT
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rfonts.set(qn(attr), FONT)


def _set_styles(doc):
    normal = doc.styles['Normal']
    _force_font(normal, size=BODY_PT, color=BLACK)
    fmt = normal.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.15
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, (size, before, after) in HEADING_FORMAT.items():
        style = doc.styles[name]
        _force_font(style, size=size, bold=True, color=BLACK)
        style.font.italic = False
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _set_page(doc):
    for section in doc.sections:
        section.left_margin = Cm(MARGINS_CM['left'])
        section.right_margin = Cm(MARGINS_CM['right'])
        section.top_margin = Cm(MARGINS_CM['top'])
        section.bottom_margin = Cm(MARGINS_CM['bottom'])


def _set_metadata(doc, content):
    from app.models.site_settings import SiteSettings

    settings = SiteSettings.get()
    title = content.select_one('.iprm-content-page__title')
    core = doc.core_properties
    core.title = _clean(title.get_text()) if title else ''
    core.author = settings.company_legal_name or ''
    core.subject = core.title


def _add_page_number_footer(doc):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(0)

    field = OxmlElement('w:fldSimple')
    field.set(qn('w:instr'), 'PAGE')
    run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rfonts = OxmlElement('w:rFonts')
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), FONT)
    size = OxmlElement('w:sz')
    size.set(qn('w:val'), '20')
    rpr.append(rfonts)
    rpr.append(size)
    run.append(rpr)
    field.append(run)
    para._p.append(field)


def _hide_borders(table):
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement('w:%s' % edge)
        element.set(qn('w:val'), 'none')
        element.set(qn('w:sz'), '0')
        borders.append(element)
    table._tbl.tblPr.append(borders)


# --- HTML -> Word ---------------------------------------------------------

def _clean(text):
    return re.sub(r'\s+', ' ', text).strip()


def _add_inline(paragraph, node, bold=False, italic=False):
    """Додати текст вузла у абзац зі збереженням <strong>/<em>."""
    if isinstance(node, NavigableString):
        text = re.sub(r'\s+', ' ', str(node))
        if not text.strip():
            if text and paragraph.runs:
                paragraph.add_run(' ')
            return
        run = paragraph.add_run(text)
        # Вмикаємо накреслення лише коли воно потрібне: присвоєння False
        # перебило б стиль абзацу (заголовки втратили б жирність).
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        return

    if node.name in ('script', 'style', 'button', 'svg'):
        return
    is_bold = bold or node.name in ('strong', 'b')
    is_italic = italic or node.name in ('em', 'i')
    for child in node.children:
        _add_inline(paragraph, child, is_bold, is_italic)


def _add_paragraph(doc, node, style=None, prefix=None, indent=None,
                   hanging=None, space_after=None, align=None):
    para = doc.add_paragraph(style=style)
    if prefix:
        para.add_run(prefix)
    for child in node.children:
        _add_inline(para, child)
    fmt = para.paragraph_format
    if indent is not None:
        fmt.left_indent = Cm(indent)
    if hanging is not None:
        fmt.first_line_indent = Cm(-hanging)
    if space_after is not None:
        fmt.space_after = Pt(space_after)
    if align is not None:
        fmt.alignment = align
    return para


def _add_list(doc, node):
    """Списки формуємо вручну: у бланку немає стилів списків і нумерації."""
    items = node.find_all('li', recursive=False)
    ordered = node.name == 'ol'
    for index, item in enumerate(items, start=1):
        prefix = '%d) ' % index if ordered else '– '
        _add_paragraph(doc, item, prefix=prefix, indent=1.0, hanging=0.5,
                       space_after=2)


def _add_definitions(doc, node):
    for child in node.find_all(['dt', 'dd'], recursive=False):
        if child.name == 'dt':
            para = _add_paragraph(doc, child, space_after=0,
                                  align=WD_ALIGN_PARAGRAPH.LEFT)
            for run in para.runs:
                run.bold = True
        else:
            _add_paragraph(doc, child, indent=1.0, space_after=6)


def _add_table(doc, node):
    rows = node.find_all('tr')
    if not rows:
        return
    columns = max(len(row.find_all(['td', 'th'])) for row in rows)
    table = doc.add_table(rows=0, cols=columns)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for source_row in rows:
        cells = source_row.find_all(['td', 'th'])
        target = table.add_row().cells
        for index, cell in enumerate(cells[:columns]):
            para = target[index].paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for child in cell.children:
                _add_inline(para, child)
            if index == 0 or cell.name == 'th':
                for run in para.runs:
                    run.bold = True

    if columns == 2:
        for row in table.rows:
            row.cells[0].width = Cm(6.0)
            row.cells[1].width = Cm(11.0)
    doc.add_paragraph()


def _emit(doc, node):
    """Перекласти один блоковий вузол HTML у Word-розмітку."""
    if isinstance(node, NavigableString):
        return
    name = node.name
    if name in HEADING_STYLES:
        _add_paragraph(doc, node, style=HEADING_STYLES[name])
    elif name == 'p':
        _add_paragraph(doc, node)
    elif name in ('ul', 'ol'):
        _add_list(doc, node)
    elif name == 'dl':
        _add_definitions(doc, node)
    elif name == 'table':
        _add_table(doc, node)
    elif name in ('section', 'div', 'article'):
        for child in node.children:
            _emit(doc, child)


def _build_body(doc, content):
    title = content.select_one('.iprm-content-page__title')
    if title is not None:
        heading = doc.add_heading(level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for child in title.children:
            _add_inline(heading, child)

    subtitle = content.select_one('.iprm-content-page__subtitle')
    if subtitle is not None:
        _add_paragraph(doc, subtitle, align=WD_ALIGN_PARAGRAPH.CENTER,
                       space_after=14)

    for node in content.select('.legal-page__preamble p'):
        para = _add_paragraph(doc, node, space_after=12)
        para.paragraph_format.first_line_indent = Cm(1.0)

    handled = {'.iprm-content-page__title', '.iprm-content-page__subtitle',
               '.legal-page__preamble'}
    sections = content.select('section.legal-section')
    if sections:
        for section in sections:
            _emit(doc, section)
    else:
        # запасний шлях для сторінок без section.legal-section
        for node in content.children:
            if isinstance(node, NavigableString):
                continue
            classes = {'.%s' % c for c in (node.get('class') or [])}
            if classes & handled:
                continue
            _emit(doc, node)

    closing = content.select_one('.legal-page__closing')
    if closing is not None:
        para = _add_paragraph(doc, closing)
        para.paragraph_format.space_before = Pt(10)
        for run in para.runs:
            run.bold = True


def _add_signature(doc):
    """Підписний блок: посада, печатка з підписом, прізвище."""
    from app.models.site_settings import SiteSettings

    settings = SiteSettings.get()
    signer_title = current_app.config.get('LEGAL_DOCX_SIGNER_TITLE',
                                          DEFAULT_SIGNER_TITLE)
    signer_name = current_app.config.get('LEGAL_DOCX_SIGNER_NAME',
                                         DEFAULT_SIGNER_NAME)
    with zipfile.ZipFile(_letterhead_path()) as archive:
        try:
            seal = io.BytesIO(archive.read(SEAL_MEDIA))
        except KeyError:
            raise LegalDocxError(
                'У бланку %s немає зображення печатки (%s).'
                % (_letterhead_path().name, SEAL_MEDIA)
            )

    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    company = doc.add_paragraph()
    company.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    company.paragraph_format.space_after = Pt(0)
    company.paragraph_format.keep_with_next = True
    company.add_run(settings.company_legal_name or '').bold = True

    table = doc.add_table(rows=1, cols=3)
    _hide_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cells = table.rows[0].cells
    for cell, width in zip(cells, (Cm(4.5), Cm(5.5), Cm(5.0))):
        cell.width = width
        cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    cells[0].paragraphs[0].add_run(signer_title)
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[1].paragraphs[0].add_run().add_picture(seal, width=Cm(4.5))
    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cells[2].paragraphs[0].add_run(signer_name)
