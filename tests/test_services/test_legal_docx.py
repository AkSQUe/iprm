"""Тести експорту юридичних сторінок у .docx (legal_docx_service)."""
from pathlib import Path

import pytest
from docx import Document

from app.services.legal_docx_service import (
    DEFAULT_LETTERHEAD, LEGAL_PAGES, LegalDocxError, export_page,
)

# Бланк листа з логотипом і печаткою не зберігається в репозиторії, тому без
# нього тести пропускаються (CI), а не падають.
LETTERHEAD = Path(__file__).resolve().parents[2] / DEFAULT_LETTERHEAD
requires_letterhead = pytest.mark.skipif(
    not LETTERHEAD.exists(), reason='Немає бланка листа: %s' % DEFAULT_LETTERHEAD)


class TestExportPage:
    def test_unknown_page_raises(self, app):
        with pytest.raises(LegalDocxError):
            export_page('no-such-page')

    @requires_letterhead
    def test_offer_has_expected_layout(self, app, tmp_path):
        path = export_page('offer', output_dir=str(tmp_path))
        assert path.exists()

        doc = Document(path)
        section = doc.sections[0]
        assert round(section.left_margin.cm, 1) == 3.0
        assert round(section.right_margin.cm, 1) == 1.5
        assert round(section.top_margin.cm, 1) == 3.0
        assert round(section.bottom_margin.cm, 1) == 2.0

        assert doc.styles['Normal'].font.name == 'Times New Roman'
        assert doc.styles['Normal'].font.size.pt == 12

        headings = [p.text for p in doc.paragraphs
                    if p.style.name.startswith('Heading')]
        assert headings[0].startswith('Публічна оферта')
        # усі 14 розділів оферти
        assert sum(1 for h in headings if h[:2].rstrip('.').isdigit()) >= 14

    @requires_letterhead
    def test_letterhead_header_is_inherited(self, app, tmp_path):
        doc = Document(export_page('offer', output_dir=str(tmp_path)))
        header_text = '\n'.join(p.text for p in doc.sections[0].header.paragraphs)
        assert '45871060' in header_text          # ЄДРПОУ з бланка
        assert 'plasma-regen.com' in header_text
        # логотип бланка лишається у колонтитулі
        assert any(str(part.partname).startswith('/word/media')
                   for part in doc.sections[0].header.part.related_parts.values())

    @requires_letterhead
    def test_seal_is_optional(self, app, tmp_path):
        with_seal = Document(export_page('offer', output_dir=str(tmp_path / 'a')))
        without = Document(export_page('offer', output_dir=str(tmp_path / 'b'),
                                       with_seal=False))
        assert len(with_seal.tables) == len(without.tables) + 1

    @requires_letterhead
    @pytest.mark.parametrize('page_key', sorted(LEGAL_PAGES))
    def test_every_registered_page_exports(self, app, tmp_path, page_key):
        doc = Document(export_page(page_key, output_dir=str(tmp_path / page_key)))
        assert any(p.style.name == 'Heading 1' and p.text.strip()
                   for p in doc.paragraphs)
